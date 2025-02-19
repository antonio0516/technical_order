import os
import re
import shutil
import time
import traceback
from typing import Union

import aspose.words as aw
import docx
from docxcompose.composer import Composer
from fastapi import (APIRouter, File, Form, Header, HTTPException, Request,
                     UploadFile)
from fastapi.responses import FileResponse
from pydantic import BaseModel

from domain.admin_log import AdminLog, LogType
from domain.auth.Hash import Hash
from domain.auth.JWT import JWT, JWTDecodeError
from domain.database.data_model.User import User
from domain.database.database import mongo_client
from domain.exam.student import Student
from domain.validation.InputValidator import InputValidator

router = APIRouter()


class StudentLoginRequest(BaseModel):
    student_number: str
    password: str


class ResetAllStudentRequest(BaseModel):
    start_number: int
    end_number: int


class PostStudentGradeRequest(BaseModel):
    exam_token: str
    grade: int
    type: str


@router.post("/login", status_code=200)
async def student_login(
    request: StudentLoginRequest,
):
    student = Student()

    # check if account already exists
    data = student.get_student_by_student_number(request.student_number)
    if data is None:
        raise HTTPException(status_code=401, detail="學號不存在")

    if data["password"] != request.password:
        raise HTTPException(status_code=401, detail="密碼錯誤")

    token = JWT.make_jwt_token(request.student_number, is_admin="0")

    return {"token": token}


@router.get("/", status_code=200)
def get_students(
    authorization: str = Header(None),
):
    token = None
    if authorization is not None:
        token = authorization.split(" ")[1]

    jwt_info = JWT.jwt_required(token, need_admin=True)

    student = Student()
    data = student.get_students()

    data = list(data)

    for student in data:
        student["_id"] = str(student["_id"])

    return data


@router.get("/exam_token", status_code=200)
def get_exam_token(type: str, authorization: str = Header(None)):
    token = None
    if authorization is not None:
        token = authorization.split(" ")[1]

    jwt_info = JWT.jwt_required(token, False)

    student = Student()

    id = student.get_student_by_student_number(jwt_info["account"])["_id"]

    data = student.get_student_by_id(id)

    if data is None:
        raise HTTPException(status_code=404, detail="學生不存在")

    # if data["exam_flag"]:
    #     raise HTTPException(status_code=400, detail="已經申請過考試 Token")

    # update to database
    student.update_student_exam_flag(id, type, True)

    return {"exam_token": data["exam_token"]}


@router.post("/reset", status_code=201)
def reset_all_students(
    request: ResetAllStudentRequest, authorization: str = Header(None)
):
    token = None
    if authorization is not None:
        token = authorization.split(" ")[1]

    jwt_info = JWT.jwt_required(token, need_admin=True)

    if request.start_number > request.end_number:
        raise HTTPException(status_code=400, detail="起始學號不可大於結束學號")

    try:
        with mongo_client.start_session() as session:
            with session.start_transaction():
                student = Student()
                student.delete_all_students(session=session)
                student.create_students_by_continuous_number(
                    request.start_number, request.end_number, session=session
                )
    except Exception as e:
        raise HTTPException(status_code=500, detail="重置學生帳號失敗")

    admin_log = AdminLog()
    admin_log.add_log(
        "重置所有學生帳號",
        jwt_info,
        "重置",
        LogType.CRITICAL,
    )

    return {}


@router.post("/reset/{id}", status_code=201)
def reset_student_by_id(id: str, authorization: str = Header(None)):
    token = None
    if authorization is not None:
        token = authorization.split(" ")[1]

    jwt_info = JWT.jwt_required(token, need_admin=True)

    student = Student()
    data = student.get_student_by_id(id)

    if data is None:
        raise HTTPException(status_code=404, detail="學生不存在")

    student_number = data["student_number"]

    try:
        with mongo_client.start_session() as session:
            with session.start_transaction():
                student.delete_student_by_id(id, session=session)
                student.create_student(student_number, session=session)
    except Exception as e:
        raise HTTPException(status_code=500, detail="重置學生帳號失敗")

    admin_log = AdminLog()
    admin_log.add_log(
        f"重置學號 {student_number} 的學生帳號",
        jwt_info,
        "重置",
        LogType.CRITICAL,
    )

    return {}


@router.post("/grade", status_code=201)
async def post_student_grade(
    request: PostStudentGradeRequest, authorization: str = Header(None)
):
    token = None
    if authorization is not None:
        token = authorization.split(" ")[1]

    jwt_info = JWT.jwt_required(token, need_admin=False)

    reporter_number = jwt_info["account"]

    student = Student()
    data = student.get_student_by_student_number(reporter_number)

    if data is None:
        raise HTTPException(status_code=404, detail="學生不存在")

    if data["grade"][request.type] != -1:
        raise HTTPException(status_code=400, detail="已經評分過")

    if data["exam_token"] != request.exam_token:
        raise HTTPException(status_code=400, detail="未持有合法考試 Token")

    student.update_student_grade(reporter_number, request.grade, request.type)

    admin_log = AdminLog()
    admin_log.add_log(
        f"更新學號 {reporter_number} 的 {request.type} 成績",
        jwt_info,
        "更新",
        LogType.WARNING,
    )

    return {}


@router.get("/score_docx", status_code=200)
def get_score_pdf():
    student = Student()

    try:
        data = student.get_students()
        data = list(data)

        data.sort(key=lambda x: int(x["student_number"]))
    except Exception as e:
        raise HTTPException(status_code=500, detail="無法取得學生資料")

    try:
        # clear tmp folder
        folder = "./tmp"
        for the_file in os.listdir(folder):
            if the_file == ".gitignore":
                continue
            file_path = os.path.join(folder, the_file)
            try:
                if os.path.isfile(file_path):
                    os.unlink(file_path)
            except Exception as e:
                print(e)

        word_file_count = len(data) // 16 + 1 if len(data) % 16 != 0 else 0

        timestamp = int(time.time())
        template_path = "./resource/student_score_template.docx"
        temp_path_prefix = f"./tmp/{timestamp}_student_score"

        for i in range(word_file_count):
            shutil.copy2(template_path, f"{temp_path_prefix}_{i}.docx")

            # replace "ACCOUNT_{i}" to student account
            doc = docx.Document(f"{temp_path_prefix}_{i}.docx")

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            # check ACCOUNT_{i}
                            if re.match(r"ACCOUNT_\d+", paragraph.text):
                                key = paragraph.text
                                number = int(key.split("_")[-1].strip())
                                value = ""
                                if number - 1 + i * 16 < len(data):
                                    value = data[number - 1 + i * 16]["student_number"]
                                paragraph.text = paragraph.text.replace(key, value)
                                paragraph.style.font.size = docx.shared.Pt(18)
                            # check AH1W_SCORE_{i}
                            if re.match(r"AH1W_SCORE_\d+", paragraph.text):
                                key = paragraph.text
                                number = int(key.split("_")[-1].strip())
                                value = ""
                                if number - 1 + i * 16 < len(data):
                                    value = str(
                                        data[number - 1 + i * 16]["grade"]["AH-1W"]
                                    )
                                    if value == "-1":
                                        value = "尚未登記"
                                paragraph.text = paragraph.text.replace(key, value)
                                paragraph.style.font.size = docx.shared.Pt(18)
                            # check OH58D_SCORE_{i}
                            if re.match(r"OH58D_SCORE_\d+", paragraph.text):
                                key = paragraph.text
                                number = int(key.split("_")[-1].strip())
                                value = ""
                                if number - 1 + i * 16 < len(data):
                                    value = str(
                                        data[number - 1 + i * 16]["grade"]["OH-58D"]
                                    )
                                    if value == "-1":
                                        value = "尚未登記"
                                paragraph.text = paragraph.text.replace(key, value)
                                paragraph.style.font.size = docx.shared.Pt(18)

            doc.save(f"{temp_path_prefix}_{i}.docx")

        # merge to one file
        first_doc = docx.Document(f"{temp_path_prefix}_0.docx")
        # add page break
        first_doc.add_page_break()
        composer = Composer(first_doc)
        for i in range(1, word_file_count):
            doc = docx.Document(f"{temp_path_prefix}_{i}.docx")
            if i != word_file_count - 1:
                doc.add_page_break()
            composer.append(doc)
        composer.save(f"{temp_path_prefix}.docx")

        return FileResponse(f"{temp_path_prefix}.docx", filename=f"學生成績文件.docx")
    except Exception as e:
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail="無法產生成績文件")


@router.get("/{id}", status_code=200)
def get_student_by_id(id: str, authorization: str = Header(None)):
    token = None
    if authorization is not None:
        token = authorization.split(" ")[1]

    jwt_info = JWT.jwt_required(token, need_admin=True)

    student = Student()
    data = student.get_student_by_id(id)

    if data is not None:
        data["_id"] = str(data["_id"])

    return data
